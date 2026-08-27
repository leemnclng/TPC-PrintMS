from __future__ import annotations

import math
from io import BytesIO
from statistics import mean
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document

from ..models.document_analysis import DocumentAnalysis, PageMargins
from ..models.enums import DocumentFileType
from ..utils.color_analysis import ImageColorMetrics, analyze_image_color, estimate_text_ink_coverage
from ..utils.image_processing import open_image
from ..utils.page_geometry import classify_orientation, classify_paper_size, emu_to_mm
from .base import DocumentAnalyzer, InvalidDocumentError, estimate_print_time


class DocxAnalyzer(DocumentAnalyzer):
    file_type = DocumentFileType.docx

    def analyze(self, filename: str, data: bytes, mime_type: str) -> DocumentAnalysis:
        try:
            document = Document(BytesIO(data))
        except Exception as error:
            raise InvalidDocumentError("The selected Word document is damaged or unsupported.") from error

        text_parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            text_parts.extend(cell.text for row in table.rows for cell in row.cells)
        text = "\n".join(part for part in text_parts if part).strip()
        word_count = len(text.split())
        stored_page_count = self._stored_page_count(data)
        page_count = stored_page_count or max(1, math.ceil(word_count / 500))
        section = document.sections[0]
        width_mm = emu_to_mm(section.page_width)
        height_mm = emu_to_mm(section.page_height)
        image_metrics: list[ImageColorMetrics] = []
        image_count = 0
        for part in document.part.related_parts.values():
            if not getattr(part, "content_type", "").startswith("image/"):
                continue
            image_count += 1
            try:
                image_metrics.append(analyze_image_color(open_image(part.blob)))
            except Exception:
                pass

        has_colored_text = self._has_colored_text(document)
        colored_content = any(metric.is_colored for metric in image_metrics) or has_colored_text
        color_pages = page_count if colored_content else 0
        bw_pages = page_count - color_pages
        warnings = ["DOCX page and color counts use saved layout metadata and may differ after opening in another Word version."]
        if stored_page_count is None:
            warnings.append("No saved page count was present, so pagination was estimated from document length.")
        coverage = round(mean(metric.coverage_percent for metric in image_metrics), 1) if image_metrics else 0.0
        text_ink = estimate_text_ink_coverage(len(text), page_count)
        image_ink = mean(metric.ink_coverage_percent for metric in image_metrics) if image_metrics else 0.0
        image_color = mean(metric.color_coverage_percent for metric in image_metrics) if image_metrics else 0.0
        ink = round(min(100.0, text_ink + image_ink), 1)
        color_coverage = round(min(100.0, image_color + (text_ink if has_colored_text else 0.0)), 1)
        warnings.append("Ink coverage is estimated from text density and embedded images because DOCX is not rendered.")
        return DocumentAnalysis(
            filename=filename,
            file_type=self.file_type,
            mime_type=mime_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size_bytes=len(data),
            page_count=page_count,
            paper_size=classify_paper_size(width_mm, height_mm),
            orientation=classify_orientation(width_mm, height_mm),
            width_mm=width_mm,
            height_mm=height_mm,
            dpi=None,
            character_count=len(text),
            word_count=word_count,
            ocr_required=False,
            image_count=image_count,
            contains_images=image_count > 0,
            image_coverage_percent=coverage,
            estimated_color_coverage_percent=color_coverage,
            estimated_ink_coverage_percent=ink,
            table_count=len(document.tables),
            graphic_count=image_count,
            margins=PageMargins(
                top_mm=emu_to_mm(section.top_margin),
                right_mm=emu_to_mm(section.right_margin),
                bottom_mm=emu_to_mm(section.bottom_margin),
                left_mm=emu_to_mm(section.left_margin),
            ),
            color_pages=color_pages,
            bw_pages=bw_pages,
            duplex_compatible=page_count > 1,
            estimated_print_time_seconds=estimate_print_time(color_pages, bw_pages),
            confidence=0.86 if stored_page_count else 0.72,
            warnings=warnings,
        )

    @staticmethod
    def _stored_page_count(data: bytes) -> int | None:
        try:
            with ZipFile(BytesIO(data)) as archive:
                root = ElementTree.fromstring(archive.read("docProps/app.xml"))
            pages = root.find("{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages")
            value = int(pages.text) if pages is not None and pages.text else 0
            return value or None
        except Exception:
            return None

    @staticmethod
    def _has_colored_text(document) -> bool:
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                rgb = run.font.color.rgb
                if rgb and max(rgb) - min(rgb) > 14:
                    return True
        return False
